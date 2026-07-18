"""Emploi core logic — UI-free so it can be tested end-to-end without Streamlit."""

import io
import json
import os
import re
from datetime import datetime

PROFILE_KEYS = ["name", "title", "location", "experience", "skills", "education", "goals"]

# ---------------- Billing tiers ----------------
# Business decision (Joy, 2026-07-14): Free/Pro/Max, priced in Naira/month.
# The number here is what's actually enforced server-side; TIER_PRICES_NGN
# is display-only (Paystack plan codes are the real source of truth for
# what a subscription actually charges — these are shown in the UI and
# used to sanity-check the plan code environment vars point at the right
# amount, never used to charge anyone directly).
TIER_LIMITS = {"free": 10, "pro": 50, "max": 300}
TIER_PRICES_NGN = {"free": 0, "pro": 3500, "max": 7500}
TIER_ORDER = ["free", "pro", "max"]


def monthly_generation_limit(tier: str) -> int:
    return TIER_LIMITS.get(tier, TIER_LIMITS["free"])

SKILLS_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "skills")
_skill_cache = {}


def load_skill(name: str) -> str:
    """Load a skill file (markdown prompt module) from skills/. Cached."""
    if name not in _skill_cache:
        try:
            with open(os.path.join(SKILLS_DIR, f"{name}.md"), encoding="utf-8") as f:
                _skill_cache[name] = f.read()
        except OSError:
            _skill_cache[name] = ""
    return _skill_cache[name]


def admin_allowed(email, allowlist) -> bool:
    """Admin-console access policy for the Streamlit app.

    `allowlist` is a comma-separated email string (EMPLOI_ADMIN_EMAILS).
    Empty/None allowlist = no restriction (open, as before). When an
    allowlist is configured, only listed emails pass; a missing email fails
    closed. Comparison is case- and whitespace-insensitive.
    """
    allowed = [e.strip().lower() for e in (allowlist or "").split(",") if e.strip()]
    if not allowed:
        return True
    return (email or "").strip().lower() in allowed


# ---------------- Prompts ----------------

def _entries_or_text(value, fallback="") -> str:
    """Normalize an experience/education field for prompt injection. Accepts
    either a plain string (legacy profile schema) or a list of
    {"summary": "..."} dicts (Career Twin wizard schema, see
    build_career_twin_extraction_prompt). Falls back to `fallback`
    (e.g. bio) when nothing structured is present, rather than rendering
    a blank "None" line into the generation prompt."""
    if isinstance(value, list) and value:
        parts = []
        for item in value:
            s = str(item.get("summary", "")).strip() if isinstance(item, dict) else str(item).strip()
            if s:
                parts.append(s)
        if parts:
            return "; ".join(parts)
    if isinstance(value, str) and value.strip():
        return value.strip()
    return str(fallback or "")


def _profile_block(profile: dict) -> str:
    """Render a candidate profile for prompt injection (generation, review,
    CV, interview prep). Supports both the legacy flat profile schema
    (name/title/location/experience/skills/education/goals, all strings)
    and the Career Twin wizard schema (headline/current_role/bio/
    career_goals, skills as a list, structured experience/education
    entries) — a Career Twin dict must never silently render "None" for
    fields that exist under a different key."""
    skills = profile.get("skills", "")
    if isinstance(skills, list):
        skills = ", ".join(str(s) for s in skills)
    goals = profile.get("goals")
    if not goals:
        goals = profile.get("career_goals", "")
    if isinstance(goals, list):
        goals = ", ".join(str(g) for g in goals)
    lines = [
        f"- Name: {profile.get('name', '')}",
        f"- Title: {profile.get('title') or profile.get('headline') or profile.get('current_role', '')}",
        f"- Location: {profile.get('location', '')}",
        f"- Experience: {_entries_or_text(profile.get('experience'), profile.get('bio'))}",
        f"- Skills: {skills}",
        f"- Education: {_entries_or_text(profile.get('education'))}",
        f"- Goals: {goals}",
    ]
    return "\n".join(lines)


def build_prompt(profile: dict, job_text: str, company: str = "") -> str:
    return f"""Act as an expert career coach and recruiter. Create a highly targeted application.

You MUST follow this writing style guide:
{load_skill('writing_style')}

You MUST evaluate fit using this rubric:
{load_skill('evaluation')}

**Candidate Profile (ground truth — nothing beyond this may be claimed):**
{_profile_block(profile)}

**Job Opportunity:**
Company: {company or 'Unknown'}
Description: {job_text}

**Output (use these exact section headers):**
## Cover Letter
A 1-page cover letter following the style guide.

## CV Bullet Points
6-8 tailored bullet points that best match this role. If any bullet is a stretch per the backtrack test, append "(stretch — verify)" to it.

## Fit Evaluation
A markdown table scoring each rubric dimension (Skills, Experience, Culture/Conditions, Career Alignment) with a one-line note each, then the 2-3 biggest gaps, then the verdict.
End with a line formatted exactly as "Fit Score: NN/100" (the weighted average).

Sound professional yet human. Never invent experience."""


def build_review_prompt(profile: dict, job_text: str, draft: str) -> str:
    return f"""You are a rigorous hiring-manager reviewer. Critique and IMPROVE the draft application below.

Rules:
- Enforce this style guide strictly; rewrite anything that violates it:
{load_skill('writing_style')}
- Remove anything generic, exaggerated, or not grounded in the candidate profile.
- Tighten language; strengthen keyword alignment with the job description.
- Keep the exact same output structure (## Cover Letter, ## CV Bullet Points, ## Fit Evaluation ending with "Fit Score: NN/100").
- Re-check the fit scores against the rubric; correct them if the draft was too generous.
- Return ONLY the final improved version, no meta-commentary.

**Candidate Profile (ground truth — nothing beyond this may be claimed):**
{profile}

**Job Description:**
{job_text}

**Draft to improve:**
{draft}"""


def build_profile_extraction_prompt(cv_text: str) -> str:
    return f"""Extract a candidate profile from this CV text. Return ONLY a JSON object with exactly these keys (all string values, use "" if unknown):
{json.dumps(PROFILE_KEYS)}

Guidance:
- "experience": a rich multi-line summary of roles, companies, dates, and achievements (keep detail).
- "skills": comma-separated list.
- "education": degrees and certifications.
- "goals": infer target roles/preferences only if stated; otherwise "".
- Never invent facts not present in the CV.

CV text:
{cv_text}"""


def build_classification_prompt(text: str) -> str:
    return f"""Classify this document. Reply with EXACTLY one word:
- CV — if it is a person's resume/CV
- JOBS — if it contains one or more job postings/listings/vacancies
- OTHER — anything else

Document:
{text[:6000]}"""


def build_jobs_extraction_prompt(text: str) -> str:
    return f"""Extract every job posting from this document. Return ONLY a JSON array; each item:
{{"company": "", "title": "", "description": "", "contact": ""}}
- "description": the full posting text for that job (responsibilities, requirements, how to apply).
- "contact": the application email address or URL if present.
- Use "" when a field is unknown. Never invent jobs.

Document:
{text}"""


def _preferences_block(profile: dict) -> str:
    """Explicit candidate-preferences lines for the match prompt. Empty
    string when the profile has none (legacy profiles, tests)."""
    lines = []
    for key, label in (("remote_preference", "Work arrangement"),
                       ("preferred_locations", "Locations"),
                       ("preferred_roles", "Target roles"),
                       ("preferred_industries", "Industries"),
                       ("employment_type", "Employment type")):
        value = profile.get(key)
        if isinstance(value, list):
            value = ", ".join(str(v) for v in value if str(v).strip())
        if value and str(value).strip():
            lines.append(f"- {label}: {value}")
    if not lines:
        return ""
    return ("\nCandidate preferences (weigh these in the location/role parts "
            "of the rubric — a remote-only candidate scored against an "
            "on-site role in a country they didn't list must lose points, "
            "and a remote role restricted to a region they didn't list is "
            "a real gap the reason must name):\n" + "\n".join(lines) + "\n")


def build_match_prompt(profile: dict, jobs: list) -> str:
    listing = "\n\n".join(
        f"[{i}] {j.get('company','?')} — {j.get('title','?')}"
        + (f" ({j.get('location')}{', remote' if j.get('is_remote') else ''})"
           if j.get('location') or j.get('is_remote') else "")
        + f"\n{str(j.get('description',''))[:600]}"
        for i, j in enumerate(jobs))
    return f"""You are a recruiter. Score how well this candidate fits EACH job below,
using this rubric (fit_score = the weighted average it defines):
{load_skill('evaluation')}

Candidate profile:
{profile}
{_preferences_block(profile)}
Jobs:
{listing}

Return ONLY a JSON array, one item per job:
[{{"index": 0, "fit_score": 0-100, "reason": "one short sentence naming the biggest strength AND biggest gap"}}]"""


# ---------------- Preference gating (hard filter before LLM scoring) ----------------

_OPEN_LOCATION_RE = re.compile(r"anywhere|global|worldwide|open to", re.IGNORECASE)


def job_passes_preferences(profile: dict, job: dict) -> bool:
    """Deterministic pre-filter: should this job even be scored for this
    candidate? Errs toward inclusion — the LLM rubric still ranks what
    passes; this only removes jobs that plainly contradict the candidate's
    stated work-arrangement/location preferences.

    Rules:
    - No preferences on the profile -> everything passes (legacy profiles).
    - Remote jobs always pass for remote-inclusive candidates ("Remote",
      "Remote or Hybrid"); a remote role restricted to a region the
      candidate didn't list is left to the LLM rubric to penalize —
      substring matching can't resolve geography honestly.
    - On-site/hybrid jobs need a CONCRETE preferred-location match
      ("Nigeria" in "Lagos, Nigeria"). Wildcard preferences like
      "Anywhere in Africa"/"Global" open remote acceptance but do NOT
      make far-away on-site roles commutable, so they don't count here.
    - A candidate with no concrete locations: on-site jobs fail if they
      said remote-only/remote-or-hybrid (the job plainly contradicts the
      stated arrangement), and pass otherwise (nothing to gate on).
    """
    remote_pref = str(profile.get("remote_preference") or "").lower()
    prefs = profile.get("preferred_locations")
    prefs = [str(p).strip() for p in prefs if str(p).strip()] if isinstance(prefs, list) else []
    if not remote_pref and not prefs:
        return True

    job_remote = bool(job.get("is_remote"))
    job_location = str(job.get("location") or "").lower()
    concrete = [p for p in prefs if not _OPEN_LOCATION_RE.search(p)]
    onsite_ok = (any(p.lower() in job_location for p in concrete) if concrete
                 else "remote" not in remote_pref)

    if job_remote:
        # A remote job only contradicts an explicitly on-site-only candidate
        # with concrete locations the job doesn't mention.
        return "remote" in remote_pref or not remote_pref or onsite_ok
    return onsite_ok


def filter_jobs_by_preferences(profile: dict, jobs: list):
    """Split jobs into (kept, skipped_count) using job_passes_preferences."""
    kept = [j for j in jobs if job_passes_preferences(profile, j)]
    return kept, len(jobs) - len(kept)


def build_interview_prompt(profile: dict, job_text: str, company: str = "") -> str:
    return f"""Act as an expert interview coach. Prepare this candidate for an interview.

Follow this preparation skill exactly:
{load_skill('interview_prep')}

**Candidate Profile (ground truth — STAR examples must come from this):**
{_profile_block(profile)}

**The role they're interviewing for:**
Company: {company or 'Unknown'}
Description: {job_text}"""


def build_chat_prompt(profile: dict, question: str, history: str = "") -> str:
    return f"""You are Emploi, a friendly expert career agent talking WITH the candidate.

{load_skill('emploi_context')}

Candidate profile (fields: {", ".join(PROFILE_KEYS)}):
{profile}

Recent conversation:
{history or "(start of conversation)"}

The candidate says:
{question}

Respond with ONLY a JSON object:
{{"reply": "your concise, practical answer",
  "profile_updates": {{}}}}

Rules for profile_updates:
- If their message states or implies new profile information (career goals, target title, location/remote preference, new skills), put it in profile_updates using the field names above. Example: they say "I want senior marketing roles" -> {{"goals": "Senior marketing roles..."}}.
- Merge with what's already there; don't erase existing detail.
- If nothing should change, leave profile_updates empty.
- Never invent facts they didn't state."""


def chat_turn(model, profile: dict, question: str, history: str = ""):
    """One conversational turn. Returns (reply_text, profile_updates dict).
    Falls back gracefully if the model ignores the JSON format."""
    raw = model.generate_content(build_chat_prompt(profile, question, history)).text
    obj = _extract_json(raw, "{", "}")
    if isinstance(obj, dict) and "reply" in obj:
        updates = obj.get("profile_updates") or {}
        clean = {k: str(v) for k, v in updates.items()
                 if k in PROFILE_KEYS and str(v).strip()}
        return str(obj["reply"]), clean
    return raw, {}


def apply_chat_updates(twin: dict, updates: dict) -> dict:
    """Merge chat_turn's legacy-keyed profile_updates into a Career Twin dict.

    chat_turn emits the legacy PROFILE_KEYS schema (title/goals/skills as
    strings); the Career Twin stores headline/career_goals/skills as lists
    and experience/education as [{"summary": ...}] entries. Values are
    appended/merged, never overwritten wholesale — a chat remark must not
    erase a curated profile. Returns the same dict, mutated."""
    for key, value in (updates or {}).items():
        text = str(value).strip()
        if not text:
            continue
        if key == "title":
            twin["headline"] = text
        elif key == "goals":
            goals = twin.get("career_goals")
            goals = list(goals) if isinstance(goals, list) else ([goals] if goals else [])
            if text not in goals:
                goals.append(text)
            twin["career_goals"] = goals
        elif key == "skills":
            existing = twin.get("skills")
            existing = list(existing) if isinstance(existing, list) else normalize_skills(existing)
            lowered = {s.lower() for s in existing if isinstance(s, str)}
            for skill in normalize_skills(text):
                if skill.lower() not in lowered:
                    existing.append(skill)
                    lowered.add(skill.lower())
            twin["skills"] = existing
        elif key in ("experience", "education"):
            entries = twin.get(key)
            entries = list(entries) if isinstance(entries, list) else []
            entries.append({"summary": text})
            twin[key] = normalize_entries(entries)
        elif key in ("name", "location"):
            twin[key] = text
    return twin


# ---------------- Parsing ----------------

FIT_RE = re.compile(r"fit\s*score[:\s]*\**\s*(\d{1,3})\s*/\s*100", re.IGNORECASE)


def parse_fit_score(text: str):
    m = FIT_RE.search(text or "")
    if m:
        score = int(m.group(1))
        return score if 0 <= score <= 100 else None
    return None


# ONE definition of an evaluation header, used by BOTH split_application and
# strip_evaluation. build_prompt asks for "## Fit Evaluation", but real output
# also uses "## Fit Score" (see the canned fixture in test_e2e). If these two
# functions disagree on what an evaluation looks like, the section they miss
# leaks into an exported artifact — which is the whole thing we're preventing.
_EVAL_HEADER = r"^\s{0,3}#{1,6}\s*\**\s*fit\s+(?:evaluation|score)\b"
_EVAL_LINE_RE = re.compile(_EVAL_HEADER, re.IGNORECASE)
_EVAL_BLOCK_RE = re.compile(_EVAL_HEADER, re.IGNORECASE | re.MULTILINE)

# Section headers emitted by build_prompt. Matched loosely (any heading level,
# optional bold/punctuation) because model output drifts; the parser must never
# raise and must degrade to "everything is the cover letter" rather than lose
# the user's draft.
_SECTION_PATTERNS = (
    ("cover_letter", re.compile(r"^\s{0,3}#{1,6}\s*\**\s*cover\s+letter\b", re.IGNORECASE)),
    ("cv_bullets", re.compile(r"^\s{0,3}#{1,6}\s*\**\s*cv\s+bullet", re.IGNORECASE)),
    ("evaluation", _EVAL_LINE_RE),
)


def split_application(text: str) -> dict:
    """Split a generated application into its parts.

    Returns {"cover_letter": str, "cv_bullets": str, "evaluation": str}.

    This exists so the EVALUATION never lands in a file the candidate sends to
    an employer — it contains their own gap analysis and "(stretch — verify)"
    markers. Only cover_letter / cv_bullets are ever exported; evaluation is
    screen-only.

    Defensive by contract (same posture as the JSON parsers): unknown or
    missing headers never raise. If no header is recognised at all, the whole
    text becomes cover_letter so the draft is never silently dropped.
    """
    parts = {"cover_letter": "", "cv_bullets": "", "evaluation": ""}
    current = None
    buckets = {"cover_letter": [], "cv_bullets": [], "evaluation": []}

    for line in (text or "").split("\n"):
        matched = None
        for key, pattern in _SECTION_PATTERNS:
            if pattern.match(line):
                matched = key
                break
        if matched:
            current = matched
            continue  # drop the header line itself; callers re-title
        # Anything before the first header (a preamble, or a whole
        # headerless draft) belongs to the cover letter — never dropped.
        buckets[current or "cover_letter"].append(line)

    for key in parts:
        parts[key] = "\n".join(buckets[key]).strip()
    return parts


def _extract_json(text: str, opener: str, closer: str):
    s = (text or "").strip()
    start = s.find(opener)
    end = s.rfind(closer)
    if start == -1 or end == -1 or end <= start:
        return None
    try:
        return json.loads(s[start:end + 1])
    except json.JSONDecodeError:
        return None


def parse_profile_json(text: str) -> dict:
    """Parse the model's profile JSON, tolerating ```json fences and extra prose."""
    raw = _extract_json(text, "{", "}")
    if not isinstance(raw, dict):
        return {}
    return {k: str(raw.get(k, "") or "") for k in PROFILE_KEYS}


def parse_json_array(text: str) -> list:
    raw = _extract_json(text, "[", "]")
    return raw if isinstance(raw, list) else []


def strip_evaluation(text: str) -> str:
    """Remove the Fit Evaluation section from an application — downloads must
    contain only sendable content; the evaluation is for on-screen reading.

    Shares _EVAL_HEADER with split_application so the two can never disagree
    about what an evaluation header looks like."""
    m = _EVAL_BLOCK_RE.search(text or "")
    return text[:m.start()].rstrip() if m else (text or "")


# ---------------- CV ingestion ----------------

def pdf_to_text(data: bytes) -> str:
    from pypdf import PdfReader
    reader = PdfReader(io.BytesIO(data))
    return "\n".join((page.extract_text() or "") for page in reader.pages)


def extract_profile(model, cv_text: str) -> dict:
    resp = model.generate_content(build_profile_extraction_prompt(cv_text))
    return parse_profile_json(resp.text)


# ---------------- Career Twin extraction (dashboard onboarding) ----------------
# The wizard's schema, NOT the legacy PROFILE_KEYS shape. String fields default
# to ""; skills is always a list. experience_years is one of EXPERIENCE_BUCKETS
# so the wizard's <select> always has a matching option.

EXPERIENCE_BUCKETS = ["1 year", "2 years", "3 years", "4 years", "5 years",
                      "6–10 years", "10+ years"]

CAREER_TWIN_TEXT_KEYS = ["name", "headline", "current_role", "location", "bio"]


MAX_TWIN_ENTRIES = 15  # defensive cap on experience/education list length


def build_career_twin_extraction_prompt(cv_text: str) -> str:
    return f"""Extract a candidate profile from this CV text. Return ONLY a JSON object with exactly these keys:
{{"name": "", "headline": "", "current_role": "", "experience_years": 0, "location": "", "skills": [], "bio": "", "experience": [], "education": []}}

Guidance:
- "headline": the candidate's professional identity in a few words (e.g. "Product Designer", "Backend Engineer").
- "current_role": their most recent role and company (e.g. "Designer at Paystack").
- "experience_years": total years of professional experience as an integer.
- "skills": JSON array of individual skill strings.
- "bio": 2-3 sentence first-person professional summary grounded ONLY in the CV.
- "experience": JSON array of past roles, most recent first. Each item is
  {{"summary": "one line: role, company, dates, and the single biggest
  achievement or responsibility"}}. One item per distinct role.
- "education": JSON array of qualifications. Each item is
  {{"summary": "one line: degree/certification, institution, year"}}.
- Use "" (or 0 / []) when unknown. Never invent facts not present in the CV.

CV text:
{cv_text}"""


def normalize_experience_years(value) -> str:
    """Map an int/str year count onto the wizard's bucket labels."""
    try:
        years = int(float(str(value).strip().split()[0].rstrip("+")))
    except (ValueError, IndexError):
        return ""
    if years <= 0:
        return ""
    if years <= 5:
        return "1 year" if years == 1 else f"{years} years"
    if years <= 10:
        return "6–10 years"
    return "10+ years"


def normalize_skills(value) -> list:
    """Skills as a clean list of strings, whether the model sent a list or a
    comma/semicolon-separated string."""
    if isinstance(value, list):
        items = value
    elif isinstance(value, str):
        items = re.split(r"[,;]", value)
    else:
        return []
    return [s for s in (str(i).strip() for i in items) if s]


def normalize_entries(value) -> list:
    """Normalize an experience/education field into [{"summary": "..."}],
    whatever shape the model sent it in (list of dicts, list of strings, or
    a single string). Drops empty entries; caps length defensively."""
    if isinstance(value, str):
        value = [value] if value.strip() else []
    if not isinstance(value, list):
        return []
    out = []
    for item in value[:MAX_TWIN_ENTRIES]:
        s = str(item.get("summary", "")).strip() if isinstance(item, dict) else str(item).strip()
        if s:
            out.append({"summary": s})
    return out


def parse_career_twin_json(text: str) -> dict:
    """Parse and normalize the model's career-twin JSON. Returns {} on garbage;
    every field is type-safe for the wizard (strings, list, bucket label)."""
    raw = _extract_json(text, "{", "}")
    if not isinstance(raw, dict):
        return {}
    twin = {k: str(raw.get(k, "") or "").strip() for k in CAREER_TWIN_TEXT_KEYS}
    twin["skills"] = normalize_skills(raw.get("skills"))
    twin["experience_years"] = normalize_experience_years(raw.get("experience_years"))
    twin["experience"] = normalize_entries(raw.get("experience"))
    twin["education"] = normalize_entries(raw.get("education"))
    # A JSON object with nothing usable in it is a failed extraction, not a twin.
    return twin if any(twin.values()) else {}


def extract_career_twin(model, cv_text: str) -> dict:
    """CV text → wizard-schema Career Twin dict. Empty dict when the model
    output is unusable (caller decides how to surface that)."""
    resp = model.generate_content(build_career_twin_extraction_prompt(cv_text))
    return parse_career_twin_json(resp.text)


def classify_document(model, text: str) -> str:
    """Return 'cv', 'jobs', or 'other'."""
    resp = model.generate_content(build_classification_prompt(text))
    word = (resp.text or "").strip().upper()
    if "CV" in word.split() or word.startswith("CV"):
        return "cv"
    if "JOBS" in word:
        return "jobs"
    return "other"


def extract_jobs(model, text: str) -> list:
    """Extract job postings from document text -> [{company, title, description}]."""
    resp = model.generate_content(build_jobs_extraction_prompt(text))
    jobs = []
    for j in parse_json_array(resp.text):
        if isinstance(j, dict) and str(j.get("description", "")).strip():
            jobs.append({"company": str(j.get("company", "") or ""),
                         "title": str(j.get("title", "") or ""),
                         "description": str(j.get("description", "") or ""),
                         "contact": str(j.get("contact", "") or "")})
    return jobs


def match_jobs(model, profile: dict, jobs: list) -> list:
    """Score every job against the profile in ONE model call.
    Returns jobs annotated with fit_score/reason, ranked best-first."""
    scores = {int(s["index"]): s for s in parse_json_array(
        model.generate_content(build_match_prompt(profile, jobs)).text)
        if isinstance(s, dict) and "index" in s}
    ranked = []
    for i, j in enumerate(jobs):
        s = scores.get(i, {})
        fit = s.get("fit_score")
        ranked.append({**j, "index": i,
                       "fit_score": int(fit) if isinstance(fit, (int, float)) else None,
                       "reason": str(s.get("reason", "") or "")})
    ranked.sort(key=lambda r: (r["fit_score"] is None, -(r["fit_score"] or 0)))
    return ranked


def resolve_job(target: str, jobs: list):
    """Find a job by 1-based number ('2', '#2') or company/title substring."""
    t = (target or "").strip().lstrip("#").strip()
    if t.isdigit():
        i = int(t) - 1
        return jobs[i] if 0 <= i < len(jobs) else None
    low = t.lower()
    for j in jobs:
        hay = f"{j.get('company','')} {j.get('title','')}".lower()
        if low and low in hay:
            return j
    return None


# ---------------- Intent routing (the "agent" brain, rule-based v1) ----------------

JD_HINTS = ("responsibilit", "requirement", "qualification", "we are looking",
            "we're looking", "about the role", "job description", "apply",
            "what you'll do", "who you are", "benefits", "salary")

BATCH_RE = re.compile(r"^\s*batch\s*(\d+)?", re.IGNORECASE)
APPLY_RE = re.compile(r"^\s*apply(?:\s+to)?\s+(.+)$", re.IGNORECASE)


def detect_intent(text: str = "", has_pdf: bool = False, has_sheet: bool = False):
    """Return (intent, arg).
    Intents: process_pdf, import_jobs, apply, batch, match, tracker, generate, chat."""
    if has_pdf:
        return "process_pdf", None
    if has_sheet:
        return "import_jobs", None
    t = (text or "").strip()
    if t.startswith("/"):
        t = t[1:].lstrip()
    m = APPLY_RE.match(t)
    if m and len(t) < 80:
        return "apply", m.group(1).strip()
    m = re.match(r"^\s*(?:interview|prep(?:are)?)(?:\s+(?:me\s+)?(?:for\s+)?(.*))?$",
                 t, re.IGNORECASE)
    if m and len(t) < 80:
        return "interview", (m.group(1) or "").strip() or None
    m = re.match(r"^\s*verify\s+(.+)$", t, re.IGNORECASE)
    if m and len(t) < 120:
        return "verify", m.group(1).strip()
    m = BATCH_RE.match(t)
    if m:
        return "batch", int(m.group(1)) if m.group(1) else 5
    if re.search(r"\b(tracker|application status|show applications)\b", t, re.IGNORECASE):
        return "tracker", None
    if len(t) < 80 and re.search(r"\b(match(es)?|rank|best fit|find.*(job|match))\b", t, re.IGNORECASE):
        return "match", None
    low = t.lower()
    if len(t) > 400 or sum(h in low for h in JD_HINTS) >= 2:
        return "generate", None
    return "chat", None


# ---------------- Generation ----------------

def generate_application(model, profile: dict, job_text: str, company: str = "",
                         review: bool = True) -> dict:
    """Generate a tailored application. `model` is any object with
    .generate_content(prompt) returning an object with .text (real Gemini or a fake)."""
    draft = model.generate_content(build_prompt(profile, job_text, company)).text
    final = draft
    if review:
        final = model.generate_content(
            build_review_prompt(profile, job_text, draft)
        ).text
    return {
        "date": datetime.now().strftime("%Y-%m-%d"),
        "company": company or "Unnamed Company",
        "result": final,
        "draft": draft,
        "fit_score": parse_fit_score(final),
        "reviewed": review,
    }


def batch_generate(model, profile: dict, jobs: list, jd_key: str,
                   company_key: str = None, review: bool = False,
                   limit: int = 10, progress=None) -> list:
    """Generate applications for up to `limit` jobs, ranked by fit score (desc)."""
    results = []
    subset = jobs[:limit]
    for i, job in enumerate(subset):
        jd = str(job.get(jd_key, "") or "")
        if not jd.strip():
            continue
        company = str(job.get(company_key, "") or "") if company_key else ""
        results.append(generate_application(model, profile, jd, company, review))
        if progress:
            progress(i + 1, len(subset))
    results.sort(key=lambda r: (r["fit_score"] is None, -(r["fit_score"] or 0)))
    return results


def guess_columns(jobs: list):
    """Heuristically pick (jd_column, company_column) from an imported sheet."""
    if not jobs:
        return None, None
    cols = list(jobs[0].keys())
    sample = jobs[: min(10, len(jobs))]

    def avg_len(c):
        vals = [str(r.get(c, "") or "") for r in sample]
        return sum(len(v) for v in vals) / max(len(vals), 1)

    jd_col = max(cols, key=avg_len)
    company_col = None
    for c in cols:
        if c.lower() in ("co", "co.", "company", "employer", "org", "organization") \
                or re.search(r"company|employer|organi[sz]ation", c, re.IGNORECASE):
            company_col = c
            break
    return jd_col, company_col


def prepare_interview(model, profile: dict, job_text: str, company: str = "") -> str:
    return model.generate_content(
        build_interview_prompt(profile, job_text, company)).text


def build_cv_prompt(profile: dict, job_text: str, company: str = "") -> str:
    return f"""Act as an expert CV writer. Write a complete, tailored CV for this candidate targeting this specific job.

Follow this CV skill exactly:
{load_skill('cv_template')}

Also obey these style rules:
{load_skill('writing_style')}

**Candidate Profile (ground truth — the ONLY source of facts):**
{_profile_block(profile)}

**Target job:**
Company: {company or 'Unknown'}
Description: {job_text}

Return ONLY the CV in markdown. No commentary before or after."""


def generate_cv(model, profile: dict, job_text: str, company: str = "") -> str:
    return model.generate_content(build_cv_prompt(profile, job_text, company)).text


# ---------------- Employer Portal (Phase 2) ----------------
# Decisions locked with Joy 2026-07-16: role #1 free (accept-gated contact,
# 10-invite cap); roles 2+ unlock-gated (₦1,000/candidate, packs of min 5).

INVITE_CAP_FREE_ROLE = 10
UNLOCK_PRICE_NGN = 1000
MIN_UNLOCK_PACK = 5

_ATS_URL_PATTERNS = [
    # (host suffix, path regex, ats name, api url template)
    ("boards.greenhouse.io", r"^/([^/]+)/jobs/(\d+)", "greenhouse",
     "https://boards-api.greenhouse.io/v1/boards/{0}/jobs/{1}"),
    ("job-boards.greenhouse.io", r"^/([^/]+)/jobs/(\d+)", "greenhouse",
     "https://boards-api.greenhouse.io/v1/boards/{0}/jobs/{1}"),
    ("job-boards.eu.greenhouse.io", r"^/([^/]+)/jobs/(\d+)", "greenhouse",
     "https://boards-api.eu.greenhouse.io/v1/boards/{0}/jobs/{1}"),
    ("jobs.lever.co", r"^/([^/]+)/([0-9a-f-]+)", "lever",
     "https://api.lever.co/v0/postings/{0}/{1}?mode=json"),
    ("jobs.ashbyhq.com", r"^/([^/]+)/([^/?#]+)", "ashby",
     "https://api.ashbyhq.com/posting-public/apiPostings/{0}/{1}"),
    ("apply.workable.com", r"^/([^/]+)/j/([^/?#]+)", "workable",
     "https://apply.workable.com/api/v3/accounts/{0}/jobs/{1}"),
    ("jobs.smartrecruiters.com", r"^/([^/]+)/([^/?#]+)", "smartrecruiters",
     "https://api.smartrecruiters.com/v1/companies/{0}/postings/{1}"),
]

# Hosts we explicitly refuse to scrape — legal/ToS decision, not a gap to fix.
_REJECTED_HOSTS = ("linkedin.com", "indeed.com", "workday.com",
                   "myworkdayjobs.com", "taleo.net")

_UNSUPPORTED_HOST_DETAIL = (
    "{host} doesn't allow us to read their job pages. Two options: "
    "(1) paste the JD text directly, or (2) on the {host} page, click "
    "'Apply on Company Site' — if that link goes to Greenhouse/Lever/Ashby/"
    "Workable/SmartRecruiters, paste THAT URL instead.")


def extract_job_from_url(url: str, fetch_fn=None):
    """Employer paste-a-URL role extraction. Dispatches on hostname to the
    supported ATS single-job public APIs and normalizes via the shared
    workers.ingest_jobs normalizers (never duplicated).

    Returns: a db.upsert_job-shaped dict on success; None for an unknown
    host or a fetch/parse failure; {"error": "unsupported_host", "detail"}
    for hosts we deliberately refuse (LinkedIn/Indeed/Workday/Taleo)."""
    from urllib.parse import urlparse
    from workers.ingest_jobs import (_fetch, normalize_greenhouse_job,
                                     normalize_lever_posting,
                                     normalize_ashby_posting,
                                     normalize_workable_job,
                                     normalize_smartrecruiters_posting)
    fetch_fn = fetch_fn or _fetch
    try:
        parsed = urlparse(url if "://" in (url or "") else f"https://{url or ''}")
        host = (parsed.hostname or "").lower()
        if host.startswith("www."):
            host = host[4:]
    except Exception:
        return None
    if not host:
        return None

    for rejected in _REJECTED_HOSTS:
        if host == rejected or host.endswith("." + rejected):
            return {"error": "unsupported_host",
                    "detail": _UNSUPPORTED_HOST_DETAIL.format(host=rejected)}

    for suffix, path_re, ats, api_template in _ATS_URL_PATTERNS:
        if host != suffix:
            continue
        m = re.match(path_re, parsed.path or "")
        if not m:
            return None
        token, job_ref = m.group(1), m.group(2)
        data = fetch_fn(api_template.format(token, job_ref))
        if data is None:
            return None
        company_name = token.replace("-", " ").title()
        try:
            if ats == "greenhouse":
                fields = normalize_greenhouse_job(data, company_name)
            elif ats == "lever":
                fields = normalize_lever_posting(data, company_name)
            elif ats == "ashby":
                # Single-posting endpoint has varied between a bare object
                # and {jobs: [...]} — support both.
                posting = data
                if isinstance(data, dict) and isinstance(data.get("jobs"), list):
                    posting = data["jobs"][0] if data["jobs"] else None
                if not isinstance(posting, dict):
                    return None
                fields = normalize_ashby_posting(posting, company_name)
            elif ats == "workable":
                fields = normalize_workable_job(data, company_name)
                if fields is None:
                    return None  # unpublished job
            else:
                fields = normalize_smartrecruiters_posting(data, company_name)
        except Exception:
            return None
        if not (fields.get("title") or "").strip():
            return None
        fields["source_ats"] = ats
        fields["source_url"] = url
        return fields
    return None


def extract_single_job(model, text: str):
    """Gemini-backed extraction of ONE job from pasted JD text. Wraps the
    extract_jobs primitive; defensive parsing throughout — returns None for
    garbage input, never raises. is_remote is a regex heuristic over the
    extracted text (shared with the ingest worker)."""
    if not (text or "").strip():
        return None
    try:
        jobs = extract_jobs(model, text)
    except Exception:
        return None
    if not jobs:
        return None
    job = jobs[0]
    from workers.ingest_jobs import _is_remote
    combined = f"{job.get('title', '')} {job.get('description', '')}"
    return {
        "title": job.get("title", ""),
        "company_name": job.get("company", ""),
        "description": job.get("description", ""),
        "location": "",
        "is_remote": _is_remote(combined),
        "salary_text": None,
        "contact": job.get("contact", ""),
        "source_ats": "raw",
    }


def build_role_shortlist_prompt(role: dict, candidates: list,
                                refinement_note: str = "") -> str:
    """Role-anchored ranking prompt: one role, many Career Twin summaries.
    Mirrors build_match_prompt's JSON-array contract (fit_score 0-100 per the
    evaluation rubric)."""
    listing = "\n\n".join(
        f"[{i}] {_profile_block(c.get('twin', c))}"
        for i, c in enumerate(candidates))
    note = ""
    if (refinement_note or "").strip():
        note = ("\nThe employer reviewed a previous shortlist and asked for "
                f"this refinement — weight it heavily:\n{refinement_note.strip()}\n")
    return f"""You are a recruiter. Score how well EACH candidate below fits this ONE role,
using this rubric (fit_score = the weighted average it defines):
{load_skill('evaluation')}

The role:
Title: {role.get('title', '')}
Location: {role.get('location') or 'Unspecified'}{' (remote)' if role.get('is_remote') else ''}
Description: {str(role.get('description', ''))[:4000]}
{note}
Candidates:
{listing}

Return ONLY a JSON array, one item per candidate:
[{{"index": 0, "fit_score": 0-100, "reason": "one short sentence naming the biggest strength AND biggest gap"}}]"""


def rank_candidates_for_role(model, role: dict, candidates: list,
                             refinement_note: str = "") -> list:
    """Score opted-in candidates against one role in a single model call.
    candidates: [{user_id, twin}]. Returns [{candidate_user_id, fit_score,
    reason}] ranked best-first; unscored candidates sort last."""
    if not candidates:
        return []
    scores = {int(s["index"]): s for s in parse_json_array(
        model.generate_content(
            build_role_shortlist_prompt(role, candidates, refinement_note)).text)
        if isinstance(s, dict) and "index" in s}
    ranked = []
    for i, c in enumerate(candidates):
        s = scores.get(i, {})
        fit = s.get("fit_score")
        ranked.append({"candidate_user_id": c.get("user_id"),
                       "fit_score": int(fit) if isinstance(fit, (int, float)) else None,
                       "reason": str(s.get("reason", "") or "")})
    ranked.sort(key=lambda r: (r["fit_score"] is None, -(r["fit_score"] or 0)))
    return ranked


def invite_gate(is_free_role: bool, invites_sent: int,
                candidate_unlocked: bool) -> tuple:
    """Deterministic rule for whether an employer may send one more invite.
    Free role: hard cap of INVITE_CAP_FREE_ROLE invites, no unlock needed.
    Paid role: the candidate must have been unlocked (1 credit). Returns
    (allowed: bool, reason: str)."""
    if is_free_role:
        if invites_sent >= INVITE_CAP_FREE_ROLE:
            return (False,
                    f"Your free role includes {INVITE_CAP_FREE_ROLE} invites and "
                    f"you've used them all. This is a reasonable-use limit — "
                    "contact hello@emploihq.com if you need more.")
        return (True, "")
    if not candidate_unlocked:
        return (False,
                "Unlock this candidate first — on paid roles each invite uses "
                f"one unlock credit (₦{UNLOCK_PRICE_NGN:,} per candidate, "
                f"packs start at {MIN_UNLOCK_PACK}).")
    return (True, "")


_NOTE_SANITIZE_RE = re.compile(r"[\r\t]|\n{2,}")


def format_invite_email(invite: dict, role: dict, employer: dict,
                        candidate: dict) -> tuple:
    """(subject, body) for the invite notification. Plain text (Brevo
    textContent). The employer-written note is user-generated content:
    control characters are stripped and every line is quoted so it can't
    impersonate Emploi copy."""
    company = employer.get("company_name", "an employer")
    title = role.get("title", "a role")
    location = ("Remote" if role.get("is_remote")
                else (role.get("location") or "Location unspecified"))
    trust_level = (employer.get("trust_level") or "").lower()
    vouched = bool(employer.get("warm_intro_by"))
    if vouched or trust_level == "high":
        trust_line = "Verified employer ✅"
    elif trust_level == "medium":
        trust_line = "Employer trust: medium"
    elif trust_level == "low":
        trust_line = ("Employer trust: LOW — verify this employer before "
                      "responding, and never pay a fee or share bank/ID details.")
    else:
        trust_line = "Employer trust: not yet verified"
    subject = f"Interview invite: {title} at {company}"
    lines = [
        f"Hi {candidate.get('name') or 'there'},",
        "",
        f"{company} looked at your Career Twin and wants to interview you for:",
        f"  {title} — {location}",
        f"  {trust_line}",
    ]
    note = str(invite.get("invite_note") or "").strip()
    if note:
        note = _NOTE_SANITIZE_RE.sub(" ", note)[:500]
        lines += ["", "Their message:"]
        lines += [f"> {line}" for line in note.split("\n") if line.strip()]
    lines += [
        "",
        f"Review and respond here: https://app.emploihq.com/invites/{invite.get('id', '')}",
        "",
        "This invite expires in 14 days. You're in control — your contact "
        "details are only shared per your visibility settings.",
        "",
        "— Emploi",
    ]
    return subject, "\n".join(lines)


_CONTACT_VIEW_FIELDS = ("name", "email", "phone", "headline", "location")
_CONTACT_VIEW_LISTS = ("skills", "experience", "education", "career_goals")


def format_employer_contact_view(candidate_twin: dict) -> dict:
    """The structured Twin view an employer gets AFTER contact is unlocked
    (accept on the free role; paid unlock on paid roles). Explicitly excludes
    the raw CV, chat history, and any application history — Emploi is the
    curation layer, not a résumé passthrough. Missing scalars render as ""
    (never None); list fields as []."""
    twin = candidate_twin if isinstance(candidate_twin, dict) else {}
    out = {}
    for key in _CONTACT_VIEW_FIELDS:
        value = twin.get(key)
        if key == "headline" and not value:
            value = twin.get("title") or twin.get("current_role")
        out[key] = str(value).strip() if value else ""
    for key in _CONTACT_VIEW_LISTS:
        value = twin.get(key)
        if key == "career_goals" and not value:
            value = twin.get("goals")
        if isinstance(value, str):
            value = [value] if value.strip() else []
        if not isinstance(value, list):
            value = []
        out[key] = [
            (str(v.get("summary", "")).strip() if isinstance(v, dict) else str(v).strip())
            for v in value
            if (str(v.get("summary", "")).strip() if isinstance(v, dict) else str(v).strip())
        ]
    return out


# ---------------- Exports ----------------

def _pdf_safe(s: str) -> str:
    """Make text safe for fpdf's latin-1 core fonts without losing meaning."""
    for a, b in (("₦", "NGN "), ("—", "-"), ("–", "-"), ("’", "'"), ("‘", "'"),
                 ("“", '"'), ("”", '"'), ("•", "-"), ("…", "...")):
        s = s.replace(a, b)
    return s.encode("latin-1", "replace").decode("latin-1")


_TABLE_SEP_RE = re.compile(r"^:?-{2,}:?$")


def _parse_md_line(raw: str):
    """Classify a markdown line -> (kind, text). Kinds: blank, h1, h2, bullet, table, text."""
    s = raw.strip()
    if not s:
        return "blank", ""
    if s.startswith("#"):
        level = len(s) - len(s.lstrip("#"))
        return ("h1" if level <= 2 else "h2"), s.lstrip("#").strip().replace("**", "")
    if s.startswith("|"):
        cells = [c.strip().replace("**", "") for c in s.strip("|").split("|")]
        if all(_TABLE_SEP_RE.fullmatch(c) for c in cells if c):
            return "blank", ""
        return "table", "   ".join(c for c in cells if c)
    if s.startswith(("- ", "* ", "• ")) or re.match(r"^\*\s+", s):
        return "bullet", re.sub(r"^[-*•]\s+", "", s).replace("**", "")
    return "text", s.replace("**", "")


def make_pdf(text: str) -> bytes:
    from fpdf import FPDF
    pdf = FPDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.set_font("Helvetica", size=11)
    for raw in text.split("\n"):
        kind, s = _parse_md_line(raw)
        if kind == "blank":
            pdf.ln(4)
        elif kind in ("h1", "h2"):
            pdf.ln(2)
            pdf.set_font("Helvetica", "B", 14 if kind == "h1" else 12)
            pdf.multi_cell(0, 8, _pdf_safe(s), new_x="LMARGIN", new_y="NEXT")
            pdf.set_font("Helvetica", size=11)
        elif kind == "bullet":
            pdf.multi_cell(0, 6, _pdf_safe("  - " + s), new_x="LMARGIN", new_y="NEXT")
        elif kind == "table":
            pdf.set_font("Helvetica", size=10)
            pdf.multi_cell(0, 5.5, _pdf_safe(s), new_x="LMARGIN", new_y="NEXT")
            pdf.set_font("Helvetica", size=11)
        else:
            pdf.multi_cell(0, 6, _pdf_safe(s), new_x="LMARGIN", new_y="NEXT")
    return bytes(pdf.output())


def make_docx(text: str, title: str = "Application") -> bytes:
    """Render generated markdown-ish output into a simple formatted .docx."""
    from docx import Document
    doc = Document()
    doc.add_heading(title, level=0)
    for line in text.split("\n"):
        kind, s = _parse_md_line(line)
        if kind == "blank":
            continue
        if kind == "h1":
            doc.add_heading(s, level=1)
        elif kind == "h2":
            doc.add_heading(s, level=2)
        elif kind == "bullet":
            doc.add_paragraph(s, style="List Bullet")
        else:
            doc.add_paragraph(s)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
